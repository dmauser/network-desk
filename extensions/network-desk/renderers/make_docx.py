#!/usr/bin/env python3
"""make_docx.py - Network Desk professional whitepaper renderer.

Dependencies: python-docx + markdown-it-py only.
    pip install python-docx markdown-it-py

Usage:
    python make_docx.py --input report.md --specialist vnet-architect
    # --output OPTIONAL -> network-desk/<specialist>/reports/<input-stem>.docx
    # --outdir <dir>    overrides the base folder (default: network-desk)
    # Cover overrides:  --title --subtitle --classification --version --author

What this renderer does (and does NOT do):
  * Renders inline **bold**, *italic*, `code`, and [text](url) as REAL Word runs
    and clickable hyperlinks - in paragraphs, headings, list items, table cells,
    and blockquotes. No literal **, backticks, or []() survive in the output.
  * Builds a cover page (title + subtitle + accent rule + metadata block) sourced
    from YAML front-matter, then the --title/--subtitle flags, then the first H1.
    The --specialist value is NEVER used as the title.
  * Renders ```mermaid fences to PNG using ONLY the local Mermaid CLI (mmdc, or
    `npx -y @mermaid-js/mermaid-cli`) and embeds them as centered figures with a
    "Figure N." caption. If the CLI is unavailable it emits "[diagram omitted]",
    lists the source in an appendix, and prints the install command - it never
    dumps the source inline and never calls a hosted service.
  * Uses real Word styles (Heading 1/2/3, List Bullet/Number, Quote, Table Grid),
    a Calibri/navy theme, a TOC field (Word builds it on open), and a footer with
    a live PAGE field. Emojis are stripped; technical arrows like -> are kept.
"""
from __future__ import annotations

import argparse
import datetime
import pathlib
import sys

# Shared, dependency-free helpers (see _common.py).
sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
import _common as C  # noqa: E402


def _docx():
    try:
        from docx import Document  # noqa: F401
        from docx.shared import RGBColor, Pt, Inches  # noqa: F401
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
        from docx.oxml.ns import qn  # noqa: F401
        from docx.oxml import OxmlElement  # noqa: F401
    except ImportError:
        sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")
    import docx
    return docx


def _markdown_it():
    try:
        from markdown_it import MarkdownIt
    except ImportError:
        sys.exit("ERROR: markdown-it-py not installed. Run: pip install markdown-it-py")
    return MarkdownIt


# --- low-level OOXML helpers ------------------------------------------------
def _oxml(tag):
    from docx.oxml import OxmlElement
    return OxmlElement(tag)


def _qn(tag):
    from docx.oxml.ns import qn
    return qn(tag)


def add_toc_field(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = _oxml("w:fldChar")
    begin.set(_qn("w:fldCharType"), "begin")
    instr = _oxml("w:instrText")
    instr.set(_qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = _oxml("w:fldChar")
    sep.set(_qn("w:fldCharType"), "separate")
    placeholder = _oxml("w:r")
    pt = _oxml("w:t")
    pt.set(_qn("xml:space"), "preserve")
    pt.text = "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
    placeholder.append(pt)
    end = _oxml("w:fldChar")
    end.set(_qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)


def enable_update_fields(doc):
    """Force Word to recompute fields (TOC, PAGE) when the document is opened."""
    settings = doc.settings.element
    if settings.find(_qn("w:updateFields")) is None:
        el = _oxml("w:updateFields")
        el.set(_qn("w:val"), "true")
        settings.insert(0, el)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = _oxml("w:fldChar")
    begin.set(_qn("w:fldCharType"), "begin")
    instr = _oxml("w:instrText")
    instr.set(_qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = _oxml("w:fldChar")
    end.set(_qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._r.append(el)


def shade_cell(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = _oxml("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def add_accent_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = _oxml("w:pBdr")
    bottom = _oxml("w:bottom")
    bottom.set(_qn("w:val"), "single")
    bottom.set(_qn("w:sz"), "24")
    bottom.set(_qn("w:space"), "1")
    bottom.set(_qn("w:color"), "0E5A9C")
    pbdr.append(bottom)
    p_pr.append(pbdr)
    return p


def add_hyperlink(paragraph, url, segments):
    """Append a clickable external hyperlink (blue, underlined) carrying the
    formatting of each source segment."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = _oxml("w:hyperlink")
    hyperlink.set(_qn("r:id"), r_id)
    for seg in segments:
        text = C.strip_emoji(seg["text"])
        if not text:
            continue
        run_el = _oxml("w:r")
        rpr = _oxml("w:rPr")
        color = _oxml("w:color")
        color.set(_qn("w:val"), "1F6FB2")
        rpr.append(color)
        u = _oxml("w:u")
        u.set(_qn("w:val"), "single")
        rpr.append(u)
        if seg["bold"]:
            rpr.append(_oxml("w:b"))
        if seg["italic"]:
            rpr.append(_oxml("w:i"))
        if seg["code"]:
            rfonts = _oxml("w:rFonts")
            rfonts.set(_qn("w:ascii"), "Consolas")
            rfonts.set(_qn("w:hAnsi"), "Consolas")
            rpr.append(rfonts)
        run_el.append(rpr)
        t = _oxml("w:t")
        t.set(_qn("xml:space"), "preserve")
        t.text = text
        run_el.append(t)
        hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


# --- inline token -> Word runs ----------------------------------------------
def _seg(text, bold, italic, strike, code, href):
    return {"text": text, "bold": bold > 0, "italic": italic > 0,
            "strike": strike > 0, "code": code, "href": href}


def flatten_inline(inline_token):
    """Walk a markdown-it inline token's children into formatting-tagged segments."""
    segs = []
    bold = ital = strike = 0
    href = None
    for child in inline_token.children or []:
        ct = child.type
        if ct == "text":
            segs.append(_seg(child.content, bold, ital, strike, False, href))
        elif ct == "code_inline":
            segs.append(_seg(child.content, bold, ital, strike, True, href))
        elif ct == "strong_open":
            bold += 1
        elif ct == "strong_close":
            bold = max(0, bold - 1)
        elif ct == "em_open":
            ital += 1
        elif ct == "em_close":
            ital = max(0, ital - 1)
        elif ct == "s_open":
            strike += 1
        elif ct == "s_close":
            strike = max(0, strike - 1)
        elif ct == "link_open":
            href = child.attrs.get("href") if hasattr(child, "attrs") else child.attrGet("href")
        elif ct == "link_close":
            href = None
        elif ct == "softbreak":
            segs.append(_seg(" ", bold, ital, strike, False, href))
        elif ct == "hardbreak":
            segs.append(_seg("\n", bold, ital, strike, False, href))
        elif ct == "image":
            alt = child.content or (child.attrGet("alt") if hasattr(child, "attrGet") else "")
            if alt:
                segs.append(_seg(alt, bold, ital, strike, False, href))
        # html_inline / other: ignored
    return segs


def _add_text_run(paragraph, seg):
    text = C.strip_emoji(seg["text"])
    if not text:
        return
    run = paragraph.add_run(text)
    if seg["bold"]:
        run.bold = True
    if seg["italic"]:
        run.italic = True
    if seg["strike"]:
        run.font.strike = True
    if seg["code"]:
        run.font.name = "Consolas"


def render_inline(paragraph, inline_token):
    """Render a markdown-it inline token into real runs + hyperlinks."""
    segs = flatten_inline(inline_token)
    idx = 0
    while idx < len(segs):
        seg = segs[idx]
        if seg["href"]:
            href = seg["href"]
            group = []
            while idx < len(segs) and segs[idx]["href"] == href:
                group.append(segs[idx])
                idx += 1
            add_hyperlink(paragraph, href, group)
        else:
            _add_text_run(paragraph, seg)
            idx += 1


# --- styling ----------------------------------------------------------------
def style_document(doc):
    from docx.shared import RGBColor, Pt
    navy = RGBColor(*C.PRIMARY_RGB)
    blue = RGBColor(*C.SECONDARY_RGB)
    for name, color in (("Title", navy), ("Heading 1", navy),
                        ("Heading 2", blue), ("Heading 3", blue)):
        try:
            doc.styles[name].font.color.rgb = color
            doc.styles[name].font.name = "Calibri"
        except KeyError:
            pass
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _heading(doc, level, inline_token):
    style = f"Heading {min(level, 3)}"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph(style="Heading 1")
    render_inline(p, inline_token)
    return p


def _list_style(kind, depth):
    base = "List Number" if kind == "ordered" else "List Bullet"
    name = base if depth <= 1 else f"{base} {min(depth, 3)}"
    return name, base


def _add_list_item(doc, kind, depth, inline_token):
    name, base = _list_style(kind, depth)
    try:
        p = doc.add_paragraph(style=name)
    except KeyError:
        try:
            p = doc.add_paragraph(style=base)
        except KeyError:
            p = doc.add_paragraph()
    render_inline(p, inline_token)


def _code_block(doc, code_text):
    from docx.shared import Pt
    p = doc.add_paragraph()
    shade_paragraph(p, "F2F4F7")
    for line in C.strip_emoji(code_text).rstrip("\n").split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)


def shade_paragraph(paragraph, hex_fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = _oxml("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), hex_fill)
    p_pr.append(shd)


# --- cover page -------------------------------------------------------------
def render_cover(doc, cover):
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    navy = RGBColor(*C.PRIMARY_RGB)
    blue = RGBColor(*C.SECONDARY_RGB)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(C.strip_emoji(cover.title))
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = navy
    run.font.name = "Calibri"

    if cover.subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = sp.add_run(C.strip_emoji(cover.subtitle))
        srun.font.size = Pt(16)
        srun.font.color.rgb = blue
        srun.font.name = "Calibri"

    add_accent_rule(doc)

    rows = cover.meta_rows()
    if rows:
        doc.add_paragraph()
        for label, value in rows:
            mp = doc.add_paragraph()
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lab = mp.add_run(f"{label}:  ")
            lab.bold = True
            lab.font.size = Pt(11)
            val = mp.add_run(C.strip_emoji(str(value)))
            val.font.size = Pt(11)

    doc.add_page_break()


# --- table rendering --------------------------------------------------------
def _collect_table(tokens, start):
    """Return (rows, end_index). rows = list of (cells, is_header) where each
    cell is an inline token (or None)."""
    rows = []
    in_header = False
    current = None
    j = start + 1
    while j < len(tokens) and tokens[j].type != "table_close":
        tt = tokens[j].type
        if tt == "thead_open":
            in_header = True
        elif tt == "thead_close":
            in_header = False
        elif tt == "tr_open":
            current = {"cells": [], "header": in_header}
        elif tt == "tr_close":
            if current is not None:
                rows.append(current)
                current = None
        elif tt in ("th_open", "td_open"):
            # the next token is the cell's inline content (may be empty)
            inline = tokens[j + 1] if j + 1 < len(tokens) and tokens[j + 1].type == "inline" else None
            if current is not None:
                current["cells"].append(inline)
        j += 1
    return rows, j


def render_table(doc, rows):
    from docx.shared import RGBColor
    if not rows:
        return
    cols = max(len(r["cells"]) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.paragraphs[0].text = ""
            inline = row["cells"][c_idx] if c_idx < len(row["cells"]) else None
            if inline is not None:
                render_inline(cell.paragraphs[0], inline)
            if row["header"]:
                shade_cell(cell, "0E5A9C")
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# --- mermaid ----------------------------------------------------------------
def handle_mermaid(doc, src, ctx):
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    ctx["fig"] += 1
    n = ctx["fig"]
    stem = f"figure-{n}"
    mmd = ctx["ddir"] / f"{stem}.mmd"
    png = ctx["ddir"] / f"{stem}.png"
    ok = C.render_mermaid_png(src, mmd, png)
    if ok:
        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(str(png), width=Inches(6.3))
        caption = f"Figure {n}."
    else:
        ctx["omitted"].append((n, src))
        ctx["mermaid_missing"] = True
        caption = f"Figure {n}. [diagram omitted - local Mermaid CLI unavailable]"
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cp.add_run(caption)
    crun.italic = True
    crun.font.size = Pt(9.5)


def render_appendix(doc, omitted):
    if not omitted:
        return
    try:
        doc.add_paragraph("Appendix: Diagram Sources", style="Heading 1")
    except KeyError:
        doc.add_paragraph("Appendix: Diagram Sources")
    intro = doc.add_paragraph()
    intro.add_run(
        "The following Mermaid diagrams could not be rendered because the local "
        "Mermaid CLI was unavailable. Install it and re-run to embed them as images:"
    )
    note = doc.add_paragraph()
    note.add_run(C.INSTALL_HINT_MERMAID).font.name = "Consolas"
    for n, src in omitted:
        try:
            doc.add_paragraph(f"Figure {n}", style="Heading 3")
        except KeyError:
            doc.add_paragraph(f"Figure {n}")
        _code_block(doc, src)


# --- token walker -----------------------------------------------------------
def render_tokens(doc, tokens, ctx):
    list_stack = []  # entries: [kind]
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        tt = tok.type
        if tt == "heading_open":
            level = int(tok.tag[1])
            _heading(doc, level, tokens[i + 1])
            i += 3
            continue
        if tt == "bullet_list_open":
            list_stack.append("bullet")
            i += 1
            continue
        if tt == "ordered_list_open":
            list_stack.append("ordered")
            i += 1
            continue
        if tt in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            i += 1
            continue
        if tt == "list_item_open":
            kind = list_stack[-1] if list_stack else "bullet"
            depth = len(list_stack)
            j = i + 1
            while j < len(tokens) and tokens[j].type not in ("inline", "list_item_close",
                                                             "bullet_list_open", "ordered_list_open"):
                j += 1
            if j < len(tokens) and tokens[j].type == "inline":
                _add_list_item(doc, kind, depth, tokens[j])
                i = j + 1
                continue
            i += 1
            continue
        if tt == "paragraph_open":
            inline = tokens[i + 1]
            p = doc.add_paragraph()
            render_inline(p, inline)
            i += 3
            continue
        if tt == "blockquote_open":
            j = i + 1
            while j < len(tokens) and tokens[j].type != "blockquote_close":
                if tokens[j].type == "inline":
                    try:
                        p = doc.add_paragraph(style="Quote")
                    except KeyError:
                        p = doc.add_paragraph()
                    render_inline(p, tokens[j])
                j += 1
            i = j + 1
            continue
        if tt == "fence":
            info = (tok.info or "").strip().lower()
            if info == "mermaid":
                handle_mermaid(doc, tok.content, ctx)
            else:
                _code_block(doc, tok.content)
            i += 1
            continue
        if tt == "code_block":
            _code_block(doc, tok.content)
            i += 1
            continue
        if tt == "table_open":
            rows, end = _collect_table(tokens, i)
            render_table(doc, rows)
            i = end + 1
            continue
        if tt == "hr":
            add_accent_rule(doc)
            i += 1
            continue
        i += 1


# --- main -------------------------------------------------------------------
def main() -> int:
    ap = argparse.ArgumentParser(
        description="Render a Network Desk markdown report to a professional DOCX whitepaper.")
    ap.add_argument("--input", required=True, type=pathlib.Path)
    ap.add_argument("--output", type=pathlib.Path, default=None,
                    help="Output path. If omitted: network-desk/<specialist>/reports/<input-stem>.docx")
    ap.add_argument("--outdir", default="network-desk",
                    help="Base dir used when --output is omitted (default: network-desk)")
    ap.add_argument("--specialist", default="Network Desk")
    ap.add_argument("--title", default=None, help="Cover title override (else front-matter, else first H1)")
    ap.add_argument("--subtitle", default=None, help="Cover subtitle override")
    ap.add_argument("--classification", default=None, help="Cover classification (e.g. Confidential)")
    ap.add_argument("--version", default=None, help="Cover document version")
    ap.add_argument("--author", default=None, help="Cover author")
    args = ap.parse_args()

    docx = _docx()
    MarkdownIt = _markdown_it()

    raw = args.input.read_text(encoding="utf-8")
    meta, body = C.parse_front_matter(raw)
    body = C.decode_entities(body)
    today = datetime.date.today().isoformat()
    cover = C.resolve_cover(meta, args, body, today, args.specialist)

    md = (MarkdownIt("commonmark", {"breaks": False, "html": False})
          .enable("table").enable("strikethrough"))
    tokens = md.parse(body)

    doc = docx.Document()
    style_document(doc)

    render_cover(doc, cover)

    add_toc_field(doc)
    enable_update_fields(doc)
    doc.add_page_break()

    ctx = {
        "fig": 0,
        "omitted": [],
        "mermaid_missing": False,
        "ddir": C.diagrams_dir(args.specialist, args.outdir),
    }
    render_tokens(doc, tokens, ctx)
    render_appendix(doc, ctx["omitted"])

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = ""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"Network Desk  |  {args.specialist}  |  Page ")
    add_page_field(footer)
    footer.add_run(f"  |  {cover.date}")

    out = C.resolve_output(args.output, args.specialist, args.input.stem, "docx", args.outdir)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    if ctx["mermaid_missing"]:
        C.eprint(f"WARNING: {len(ctx['omitted'])} Mermaid diagram(s) omitted - "
                 f"local Mermaid CLI not found. Install to embed them:")
        C.eprint(f"  {C.INSTALL_HINT_MERMAID}")

    size = out.stat().st_size
    figs = ctx["fig"]
    print(f"OK  {out}  ({size:,} bytes, {figs} figure(s))")
    return 0


if __name__ == "__main__":
    sys.exit(main())
